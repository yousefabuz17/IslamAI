import asyncio
import json
import re
import threading
from collections import (OrderedDict, namedtuple)
from concurrent.futures import ThreadPoolExecutor
from configparser import ConfigParser
from copy import deepcopy
from dataclasses import dataclass
from functools import lru_cache
from multiprocessing import cpu_count
from pathlib import Path
from pprint import pprint
from random import choice
from time import time
from aiohttp import (ClientSession, TCPConnector, client_exceptions)
# from ascii_graph import Pyasciigraph
from time import sleep
from bs4 import BeautifulSoup
from docx import Document
from geocoder import location
from nested_lookup import nested_lookup as nested
from pdfminer.high_level import extract_pages
from rapidfuzz import (fuzz, process)
from tqdm import tqdm
from unidecode import unidecode
from string import ascii_lowercase
from blueprints.data_loader import DataLoader
import tracemalloc

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import (TimeoutException, WebDriverException, NoSuchElementException, NoSuchFrameException)

load, translate = DataLoader.load_file, DataLoader.translate


'''
IslamAI - Data Collection Module

This file, ai_data.py, serves as the backbone of my `IslamAI` project's data collection process.
It houses a diverse set of structured classes, each tailored to interact with specific, authentic data sources.
While these classes won't directly participate in the core functionality of my program, they play a crucial role in gathering and organizing data.

My comprehensive toolset includes asynchronous data fetching, text processing, and even advanced string matching algorithms.
Employs multi-threading for efficiency, ensuring that data retrieval remains swift and scalable.
The configuration and caching mechanisms optimize performance, while the progress tracking with tqdm keeps us informed of the process.
It signifies the first step towards building a robust and intelligent system for our users.

Notes:
    ~ It's worth noting that the inclusion of tqdm serves as a temporary debugging aid to track the process's progress.
    ~ tqdm slows the extraction process
'''

class ConfigInfo:
    _config = None
    _path = None
    
    def __init__(self, path='', key='Sources', file_name='sources'):
        config = self._get_config(path, key, file_name)
        for key, value in config.items():
            setattr(self, key, value)

    @classmethod
    @lru_cache(maxsize=1)
    def _get_config(cls, *args):
        path, key, file_name = args
        file_name = f'{file_name}.ini'
        if cls._path is None:
            cls._path = Path(__file__).parent.absolute() / path
        config_parser = ConfigParser()
        config_parser.read(cls._path / file_name)
        config = dict(config_parser[key])
        if cls._config is None:
            cls._config = config
        return config

class SingletonMeta(type):
    _instances = {}
    _lock = threading.Lock()
    
    def __call__(cls, *args, **kwargs):
        with cls._lock:
            if cls not in cls._instances:
                instance = super(SingletonMeta, cls).__call__(*args, **kwargs)
                cls._instances[cls] = instance
        return cls._instances[cls]

@dataclass
class BaseAPI(metaclass=SingletonMeta):
    '''Class for flexible methods'''
    config: ConfigInfo=ConfigInfo()
    path: Path=Path(__file__).parent.absolute() / 'islamic_data'
    
    def __init__(self):
        self.url = self.config

    @classmethod
    def get_instance(cls):
        return cls()
    
    def _get_rand_token(self):
        from random import choice
        rapid_api_url, _, config = self._get_rapidapi(g_config=True)
        tokens = [i.removeprefix('rapidapi_token') for i in list(vars(config).keys())[1:]]
        rand_choice = deepcopy(choice(tokens))
        failed_tokens = dict.fromkeys(tokens)
        new_headers = self._get_rapidapi(token_num=rand_choice)[1]
        if rand_choice not in failed_tokens:
            failed_tokens[rand_choice] = False
            return (rapid_api_url, new_headers)
        return (rapid_api_url, new_headers) if not all(list(failed_tokens.values())) else False
    
    @staticmethod
    def _get_rapidapi(token_num=1, g_config=False):
        new_config = ConfigInfo('', 'RapidAPI', 'rapidapi')
        rapid_api_url = f'https://{new_config.rapidapi_host}'
        rapid_api_token = getattr(new_config, f'rapidapi_token{token_num}')
        rapid_api_host = rapid_api_url.removeprefix('https://')
        headers = {
                "X-RapidAPI-Key": rapid_api_token,
                "X-RapidAPI-Host": rapid_api_host
            }
        if not g_config:
            return (rapid_api_url, headers)
        return (rapid_api_url, headers, new_config)
    
    async def _get_rapidapi_info(self, surahID=1):
        rapid_api_url, rapid_api_headers = self._get_rapidapi()
        rapid_api_response = await self._request(endpoint=surahID, url=rapid_api_url,
                                                headers=rapid_api_headers, slash=True,
                                                rapidapi=True)
        if not rapid_api_response:
            return self._get_rapidapi_info(surahID)
        return {
                'description':rapid_api_response.get('description', ''),
                'name_translation': rapid_api_response.get('translation', '').title()
                }
    
    async def _request(self, **kwargs):
        default_values = (self.url, '', None, False, False)
        url, endpoint, headers, slash, rapidapi = tuple(kwargs.get(key, default_values[i]) for i, key in enumerate(('url', 'endpoint', 'headers', 'slash', 'rapidapi')))
        slash = '/' if slash else ''
        full_endpoint = f'{url}{slash+str(endpoint)}'
        response = ''
        try:
            async with ClientSession(connector=TCPConnector(ssl=False, enable_cleanup_closed=True,
                                                            force_close=True, ttl_dns_cache=300),
                                                            raise_for_status=True) as session:
                async with session.get(full_endpoint, headers=headers) as response:
                    response = await response.json()
                    return response
        except (client_exceptions.ContentTypeError):
            response = await response.text()
            return response
        except (client_exceptions.ServerDisconnectedError,
                client_exceptions.ClientConnectionError):
            await asyncio.sleep(0.1)
            return await self._request(**kwargs)
        except client_exceptions.ClientResponseError as error_:
            mask_ = lambda token: token[:4]+('*'*(len(token)-35))
            rapid_api_url, new_headers = ('', '')
            if not self._get_rand_token():
                raise error_('All tokens failed.')
            if rapidapi and not response:
                await asyncio.sleep(0.1)
                headers = kwargs.get('headers', '')
                rapid_api_url, new_headers = self._get_rand_token()
                new_token = mask_(new_headers.get('X-RapidAPI-Key', ''))
            new_kwargs = {
                'url': rapid_api_url,
                'headers': new_headers
                }
            kwargs.update(new_kwargs)
            response = await self._request(**kwargs)
            if response:
                return response
            print(f"Token failed. Token ({new_token}) being used.")
    
    @staticmethod
    def best_match(string, **kwargs):
        #?> Add a check ratio method
        values_ = kwargs.get('values_', ['test1', 'test2', 'test3'])
        if values_ and not isinstance(values_, (dict, OrderedDict)):
            values_ = {value: key for value, key in enumerate(values_)}
        match_ = process.extractOne(string.lower(), [i.lower() for i in values_.values()], scorer=fuzz.ratio)
        matched = match_[0].upper() if all(i.isupper() for i in values_.values()) else match_[0].title()
        return matched, match_
    
    @staticmethod
    def _get_file(path, file_name, type_='pdfs', mode='rb', ext='pdf'):
        file = open(path / type_ / f'{file_name}.{ext}', mode=mode)
        return file
    
    @staticmethod
    def _extractor(pdf, **kwargs):
        #^ PDF Files only
        '''
        :pdf: PDF file name as BufferReader
        :kwargs: maxpages, page_numbers -> range(start,end)
        '''
        default_values = (0, None)
        maxpages, page_numbers = tuple(kwargs.get(key, default_values[i]) for i, key in enumerate(('maxpages', 'page_numbers')))
        if kwargs:
            pdf_file = extract_pages(pdf, maxpages=maxpages, page_numbers=page_numbers)
            clean_pdf = ''.join([j.get_text() for i in pdf_file for j in i if hasattr(j, 'get_text')]).split('\n')
            return clean_pdf
        else:
            pdf_file = extract_pages(pdf)
            clean_pdf = ''.join([j.get_text() for i in pdf_file for j in i if hasattr(j, 'get_text')]).split('\n')
            return clean_pdf
    
    def _get_page(self, file, start=None, end=None):
        return self._extractor(file, maxpages=start) if not end else self._extractor(file, page_numbers=range(start,end))
    
    @lru_cache(maxsize=1)
    @staticmethod
    def _load_file(path, name, mode='r', encoding='utf-8', type_='json', folder='jsons'):
        #!> Modify for flexibility
        return json.load(open(path / folder / f'{name}.{type_}', mode=mode, encoding=encoding))
    
    @classmethod
    def _exporter(cls, contents, file_name, path=''):
        for _ in tqdm(range(len(contents)), desc=f'Processing `{file_name}`', unit='MB', colour='green'):
            sleep(0.00001)
        if not path:
            path = cls.path / 'jsons' / f'{file_name}.json'
            with open(path, mode='w', encoding='utf-8') as file:
                json.dump(contents, file, indent=4, ensure_ascii=False)
        else:
            path = cls.path / path / f'{file_name}.json'
            with open(path, mode='w', encoding='utf-8') as file:
                json.dump(contents, file, indent=4, ensure_ascii=False)
        print(f'`{file_name}.json` was exported in {path}')
    
    @staticmethod
    def _get_indexes(page, pattern=None, found=False):
        norm_indexes, grouped_indexes = ([], [])
        if pattern:
            norm_indexes = [(idx, i) for idx,i in enumerate(page) if re.search(pattern, i)]
            grouped_index = [(norm_indexes[i], norm_indexes[i + 1]) for i in range(0, len(norm_indexes)-1)]
        if found:
            '''Mainly for wudu-foundations'''
            norm_indexes = [(idx, i) for idx,i in enumerate(page) if i.isupper()]
            grouped_index = [(norm_indexes[i], norm_indexes[i + 1]) for i in range(0, len(norm_indexes)-1)]
        norm_indexes.append(((norm_indexes[-1]), (len(page),norm_indexes[-1][-1])))
        grouped_index.append(((grouped_index[-1][-1]), (len(page),grouped_index[-1][-1][-1])))
        return grouped_index, norm_indexes
    
    @staticmethod
    def _get_sur_indexes(page, pattern=None, found=False):
        '''Mainly for QuranAPI'''
        if pattern:
            norm_indexes = [(idx, i) for idx,i in enumerate(page) if re.search(pattern, i)]
            grouped_index = [(norm_indexes[i], norm_indexes[i + 1]) for i in range(0, len(norm_indexes)-1)]
        if (not norm_indexes):
            norm_indexes, grouped_index = ([], [])
        else:
            norm_indexes.append(((norm_indexes[-1]), (len(page),norm_indexes[-1][-1])))
            grouped_index.append(((grouped_index[-1][-1]), (len(page),grouped_index[-1][-1][-1])))
        return grouped_index, norm_indexes
    
    async def _extract_contents(self, **kwargs):
        default_values = ['']*2+[False, self.url] +['']*3
        html_file, endpoint, slash, url, class_, tag_, style = tuple(kwargs.get(key, default_values[i]) for i,key in enumerate(('html_file', 'endpoint', 'slash', 'url', 'class_', 'tag_', 'style')))
        main_page = await self._request(url=url, endpoint=endpoint, slash=slash) if not html_file else html_file
        # try:
        #     soup = BeautifulSoup(main_page, 'html.parser')
        # except TypeError:
        #     soup = html_file
        soup = BeautifulSoup(main_page, 'html.parser')
        params = {}
        # for key, value in kwargs.items():
        #     if key not in ['html_file', 'endpoint', 'slash', 'url']:
        #         params[key] = value
        #         if key=='tag_':
        #             params['name'] = value
        #         # else:
        #         #     params['attrs'] = {key: value}
        # # params.update(kwargs)
        # # if params:
        # contents = soup.find_all(**params)
        # return contents
        # return soup
        if (class_) and (not tag_):
            params['class_'] = class_
            contents = soup.find_all(**params)
            return contents
        elif (tag_):
            params['tag_'] = tag_
            contents = soup.find_all(tag_)
            return contents
        elif (tag_ and class_):
            params['class_'] = class_
            contents = soup.find_all(tag_, **params)
            return contents
        elif (tag_ and style):
            params['style'] = class_
            contents = soup.find_all(tag_, **params)
            return contents
        elif (style):
            params['style'] = style
            contents = soup.find_all(**params)
            return contents
        return soup

    async def get_quran_keyword(self, keyword=''):
        '''Mainly for RapidAPI'''
        rapid_api_url, headers = self._get_rapidapi()
        response = await self._request(endpoint=f'corpus/{keyword}', url=rapid_api_url, slash=True, headers=headers)
        return response

    @staticmethod
    def _get_driver():
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')
        driver = webdriver.Chrome(options=options)
        return driver

    @staticmethod
    async def _get_element(*args):
        driver, by, tag_name = args
        element=None
        try:
            wait = WebDriverWait(driver, 10)
            element = wait.until(EC.presence_of_element_located((by, tag_name)))
        except (TimeoutException, WebDriverException, NoSuchElementException, NoSuchFrameException):
            if not element:
                print(f'Element `{tag_name}` not found. Trying again.')
                return await BaseAPI._get_element(*args)
        return element

    @staticmethod
    def add_line_breaks(text, words_per_line=7):
        words = re.findall(r'\S+\s*', text)
        # Adds multiple spaces every 7 words for better readability
        words_with_spaces = []
        for i, word in enumerate(words):
            words_with_spaces.append(word)
            if (i + 1) % words_per_line == 0:
                words_with_spaces.append('   ')  # Add two spaces for splitting

        description = ''.join(words_with_spaces).split('   ')
        return description

class QuranAPI(BaseAPI):
    def __init__(self):
        super().__init__()
        self.url = self.config

    async def surahquran_extract_surahs(self, export=False):
        '''lang, ayaID, surahID'''
        surahquran_endpoint = 'Surah-translation/meanings-language-{}-surah-{}.html'

        async def _get_langs():
            '''Returns {langID:lang}'''
            url = self.config.surah_quran
            main_lang_endpoint = 'Surah-translation/Translating-meanings-Quran-en.html'
            lang_soup = await self._extract_contents(url=url, endpoint=main_lang_endpoint, slash=True, class_='col-lg-12')
            unfixed_langs = [i.strip('\n') for i in [i.text.split('- ') for i in lang_soup][1]]
            chinese_to_jawa = unfixed_langs[1:unfixed_langs.index('Jawa52')+1]
            telugou = unfixed_langs[-1].split('\n')[0]
            merge_langs = chinese_to_jawa + [telugou]
            all_langs = OrderedDict({langID:re.sub(r'\d{0,2}$','',lang) for langID,lang in enumerate(merge_langs,start=1)})
            return all_langs

        async def _parse_surah(en=False, **kwargs):
            '''langID, surahID'''
            langID, surahID = tuple(kwargs.get(i) for _,i in enumerate(('langID', 'surahID')))
            url = self.url.surah_quran
            endpoint = surahquran_endpoint.format(langID, surahID)
            class_ = 'ara' if (langID==2 and not en) else 'her'
            element_contents = await self._extract_contents(url=url,endpoint=endpoint,slash=True, class_=class_)
            ara_slicer1 = slice(0,-3) if (langID==2 and not en) else slice(0,None)
            ara_slicer2 = slice(None, None, -1) if (langID==2 and not en) else slice(0,None)
            verses = [i.get_text(strip=True)[ara_slicer1][ara_slicer2] for i in element_contents]
            verse_contents = OrderedDict()
            for idx, verse in enumerate(verses, start=1):
                verseID = f'verse {surahID}:{idx}'
                if verseID not in verse_contents:
                    verse_contents[verseID] = {}
                verse_contents[verseID] = verse
            return verse_contents

        async def _parse_langs(**kwargs):
            all_langs, en_verses = await asyncio.gather(*[
                                                _get_langs(),
                                                _parse_surah(langID=2, en=True, **kwargs)])
            lang_contents = OrderedDict({'languages': {}})
            for (langID, lang) in all_langs.items():
                verses = await _parse_surah(langID=langID, **kwargs)
                lang_contents['languages'][lang] = verses
            lang_contents['languages']['English'] = en_verses
            sorted_languages = {k: v for k, v in sorted(lang_contents['languages'].items())}
            return sorted_languages

        async def _extract_all():
            surah_list = await self._surah_list()
            all_surahs = OrderedDict()
            for idx, (surahID, surah_name_) in tqdm(enumerate(surah_list.items(), start=1),
                                                    total=len(surah_list), desc='Processing Surahs (SurahQuran)',
                                                    unit='MB', colour='green'):
                surah_name, surah_name_ar = surah_name_['name_complex'], surah_name_['surah_name_ar']
                surah_contents = await _parse_langs(surahID=surahID)
                if idx not in all_surahs:
                    all_surahs[idx] = {}
                all_surahs[idx] = {
                                    'name_complex': surah_name,
                                    'surah_name_ar': surah_name_ar,
                                    'verses': {**surah_contents}
                                    }
                file_name = f'{idx}-{unidecode(surah_name)}'
                surah_base_contents = await self._surah_base_info(idx, source='`https://surahquran.com`')
                with open(self.path / 'jsons' / 'quran' / 'surah-quran' / f'{file_name}.json', mode='w', encoding='utf-8') as file:
                    full_surah = all_surahs[idx]
                    surah_base_contents.update(full_surah)
                    json.dump(surah_base_contents, file, indent=4, ensure_ascii=False)
            return all_surahs
        
        if export:
            return await self._merge_all('all-surahs-surahquran', 'surah-quran')
        all_surah_contents = await _extract_all()
        return all_surah_contents

    async def _merge_all(self, file_name, folder_name='', sorted_=True):
        all_surah_files = DataLoader(folder_path=f'jsons/quran/{folder_name}')()
        all_surahs = OrderedDict(sorted({key: value for key, value in all_surah_files.items()}.items(),key=lambda item: int(item[0].split('-')[0] if sorted_ else item)))
        self._exporter(all_surahs, file_name)
        print(f'All files for `{folder_name}` merged successfully.')
    
    async def _surah_list(self):
        #** Same contents for all surah endpoints
        altafsir_endpoint = 'ViewTranslations.asp?Display=yes&SoraNo={}&Ayah=0&toAyah=0&Language={}&LanguageID=2&TranslationBook={}'
        surahID, langID, authorID = (1,2,3)
        endpoint = altafsir_endpoint.format(*(surahID, langID, authorID))
        soup = await self._extract_contents(endpoint=endpoint, slash=True, url=self.url.altafsir, tag_='option')
        old_list = [i.get_text(strip=True).translate(str.maketrans('', '', '\xa0\r\n')).split('(') for i in soup][:114]
        updated_list = [[i[1].rstrip(')'), re.sub(r'\d{1,3}', '', i[0]).lstrip()] for i in old_list]
        surah_dict = {idx: {'name_complex': i,
                            'surah_name_ar': j}
                            for idx, (i,j) in enumerate(updated_list, start=1)}
        return surah_dict

    async def altafsir_extract_surahs(self, export=False):
        altafsir_endpoint = 'ViewTranslations.asp?Display=yes&SoraNo={}&Ayah=0&toAyah=0&Language={}&LanguageID=2&TranslationBook={}'
        #** 'lang_author_ids': {language: [langID, translator(s)ID]}
        lang_author_ids = {
                        'albanian': [27, 19], 'azerbaijani': [24, 0], 'bosnian': [19, 0],
                        'bengali': [17, 0], 'bulgarian': [28, 20], 'chinese': [8, 0],
                        'czech': [29, 0], 'dutch': [18, 0], 'english': [2, 3, 4, 5, 7, 8, 9, 10, 21], #Dont forget 18
                        'finnish': [22, 0], 'french': [3, 1, 11], 'german': [4, 0],
                        'hindi': [16, 0], 'indonesian': [13, 0], 'italian': [5, 0],
                        'japanese': [9, 0], 'korean': [25, 0], 'kurdish': [30, 0],
                        'malay': [15, 0], 'persian': [11, 0], 'portuguese': [7, 0],
                        'romanian': [20, 12, 15], 'russian': [14, 0], 'spanish': [6, 0],
                        'tamil': [23, 0], 'thai': [21, 0], 'turkish': [10, 0], 'urdu': [12, 0],
                        'uzbek': [26, 0]
                    }
        
        async def _get_lang_authors():
            lang_authors = OrderedDict({})
            #**Same langIDs for all languages
            for _, (lang, _) in enumerate(lang_author_ids.items()):
                endpoint = altafsir_endpoint.format(1, lang_author_ids.get(lang)[0], lang_author_ids.get(lang)[1])
                soup_ = await self._extract_contents(endpoint=endpoint, slash=True, url=self.url.altafsir, tag_='option')
                soup = [i.text for i in soup_]
                authors = soup[soup.index('Uzbek')+1:]
                updated_authors = [0]if not authors else list(zip(authors, lang_author_ids.get(lang)[1:]))
                lang_authors[lang] = updated_authors
            return lang_authors
        
        async def _parse_verses(surahID, langID, authorID):
            endpoint = altafsir_endpoint.format(surahID, langID, authorID)
            url = f'{self.url.altafsir}/{endpoint}'
            driver = self._get_driver()
            driver.get(url)
            iframe_element = await self._get_element(driver, By.TAG_NAME, 'iframe')
            driver.switch_to.frame(iframe_element)
            iframe_content = driver.page_source
            soup = BeautifulSoup(iframe_content, 'html.parser')
            old_contents = [i for i in ' '.join([i.text for i in soup]).split('\n') if i][1:]
            surah_rapidapi_info = await self._surah_base_info(surahID)
            verse_count = nested('verses_count', surah_rapidapi_info)[0]
            fixed_contents = ''.join(old_contents).split()
            try:
                brack_indexes = [(idx, bracket) for idx,bracket in enumerate(fixed_contents) if re.search(r'(?:\[\d{1,3}:\d{1,3}\])',bracket)]
                grouped_indexes = [(brack_indexes[i], brack_indexes[i+1]) for i in range(0, len(brack_indexes)-1)]
                grouped_indexes.append((brack_indexes[-1], (len(''.join(old_contents)), '')))
            except IndexError:
                return {'Fix': ''}
            full_verse = OrderedDict()
            for idx, (i,j) in enumerate(grouped_indexes, start=1):
                start, end = i[0], j[0]
                fix_j = re.sub(r'(?:\[(.*?)\])', '', j[1])
                org_verse = ' '.join(fixed_contents[start+1:end] + [fix_j])
                verse = ''.join(org_verse)
                #! Fix here for the remaining languages (source: altafsir)

                final_verse = ('' if not verse
                                else verse[:verse.find('*')] if re.search(r'\*', verse)
                                else (verse[:verse.find(' Sources')-5] if re.search(r'\[M\] Sources', verse)
                                else (verse[:verse.find('Sources')-4] if re.search(r'\[M\]Sources', verse)
                                else verse)))

                verseID = f'verse {surahID}:{idx}'
                if len(brack_indexes) != verse_count:
                    return {'Fix': ''}
                if verseID not in full_verse:
                    full_verse[verseID] = {}
                full_verse[f'verse {surahID}:{idx}'] = final_verse.rstrip()
            return full_verse
        
        async def _parser(**kwargs):
            #^ _parse_verses(surahID=surahID, langID=langID, authorID=authorID)
            return await _parse_verses(**kwargs)
        
        async def _parse_langs(**kwargs):
            lang_authors = await _get_lang_authors()
            all_contents = OrderedDict({'languages': {}})
            for _, (lang, author_ids) in enumerate(lang_authors.items()):
                lang_contents = {}
                for idx_, authorIDs in enumerate(author_ids, start=1):
                    author_name = None if authorIDs==0 else authorIDs[0]
                    authorID = 0 if authorIDs==0 else authorIDs[1]
                    langID = lang_author_ids.get(lang)[0]
                    author_contents = {author_name: {}}
                    author_lang_contents = await _parser(langID=langID, authorID=authorID, **kwargs)
                    author_contents[author_name] = author_lang_contents
                    lang_contents[idx_] = author_contents
                all_contents['languages'][lang] = {'translators':lang_contents}
            return all_contents
        
        async def _extract_all():
            surah_list = await self._surah_list()
            all_surahs = OrderedDict()
            for idx, (surah_id, surah_name_) in tqdm(enumerate(surah_list.items(), start=1),
                                                    total=len(surah_list), desc='Processing Surahs (Altafsir)',
                                                    unit='MB', colour='green'):
                surah_name, surah_name_ar = surah_name_['name_complex'], surah_name_['surah_name_ar']
                surah_contents = await _parse_langs(surahID=surah_id)
                if idx not in all_surahs:
                    all_surahs[idx] = {}
                all_surahs[idx] = {
                                    'name_complex': surah_name,
                                    'surah_name_ar': surah_name_ar,
                                    'verses': {**surah_contents}
                                    }
                file_name = f'{idx}-{unidecode(surah_name)}'
                surah_base_contents = await self._surah_base_info(idx, source='`https://altafsir.com`')
                with open(self.path / 'jsons' / 'quran' / 'altafsir' / f'{file_name}.json', mode='w', encoding='utf-8') as file:
                    full_surah = all_surahs[idx]
                    surah_base_contents.update(full_surah)
                    json.dump(surah_base_contents, file, indent=4, ensure_ascii=False)
            return all_surahs
        if export:
            return await self._merge_all('all-surahs-altafsir', 'altafsir')
        all_surah_contents = await _extract_all()
        return all_surah_contents

    async def extract_surahs_info(self, export=False):
        url = self.url.surah_quran
        main_endpoint = 'transliteration-aya-{}-sora-{}.html'
        
        async def _get_descr(*args):
            '''args: ayaID, surahID'''
            endpoint = main_endpoint.format(*args)
            descr = await self._extract_contents(url=url, endpoint=endpoint, slash=True, style='font-size: 12pt;')
            fix_descr = [re.sub(r'^(\*\d{1,3}\)\.?)','',i.text).rstrip('  \xa0').lstrip() for i in descr][-1]
            words_with_spaces = self.add_line_breaks(fix_descr, 7)
            return words_with_spaces

        @lru_cache(maxsize=1)
        async def _get_surah_dict():
            surah_url = self.url.surah_quran
            surah_endpoint = 'Surah-translation/Quran-language-en-6.html'
            sur_names = await self._extract_contents(url=surah_url, endpoint=surah_endpoint, slash=True, tag_='td')
            surah_names = [re.sub(r'^(\d{1,3}\- )','',i.text).strip() for i in sur_names]
            all_surahs = OrderedDict({idx: surah for idx, surah in enumerate(surah_names, start=1)})
            return all_surahs

        async def _get_surah_verses(surahID):
            surah_file = _get_surah_file()
            en = nested('English', surah_file)
            fixed = {}
            for idx,i in enumerate(en, start=1):
                surah_lst = [verse for _,verse in i.items()]
                if idx not in fixed:
                    fixed[idx] = {}
                fixed[idx] = surah_lst
            return fixed.get(surahID)
        
        @lru_cache(maxsize=1)
        async def _get_verse_count(surahID):
            verse_count = await _get_surah_verses(surahID)
            return len(verse_count)

        @lru_cache(maxsize=1)
        def _get_surah_file():
            return load('jsons', 'all-surahs-surahquran', encoding='utf-8')
        
        async def _parse_all():
            surah_file = _get_surah_file()
            surah_dict = await _get_surah_dict()
            all_surah_contents = OrderedDict()

            for surahID, surah_name in tqdm(surah_dict.items(), desc='Processing Surah Meanings', colour='green', unit='MB'):
                verse_count, surah_verses = await asyncio.gather(
                                            _get_verse_count(surahID),
                                            _get_surah_verses(surahID))
                all_verse_info = []

                for ayaID, verse in enumerate(surah_verses, start=1):
                    verse_info = OrderedDict({'verse-id': '', 'verse': '', 'description': ''})
                    verse_descr = await _get_descr(ayaID, surahID)
                    id_ = f'[{surahID}:{ayaID}]'
                    verse_info['verse-id'] = id_
                    verse_info['verse'] = verse
                    verse_info['description'] = verse_descr
                    all_verse_info.append(verse_info.copy())

                all_surah_contents[surahID] = {
                    'name': surah_name,
                    'id': surahID,
                    'verse-count': verse_count,
                    'verse-info': all_verse_info
                }
                with open(self.path / 'jsons' / 'quran' / 'verse-meanings' / f'{surahID}-{surah_name}.json', mode='w', encoding='utf-8') as file:
                    surah = all_surah_contents[surahID]
                    json.dump(surah, file, indent=4)
            return all_surah_contents
        if export:
            return await self._merge_all('all-surah-meanings', 'verse-meanings')
        all_surah_meanings = await _parse_all()
        return all_surah_meanings


    async def _surah_base_info(self, surahID=1, source=''):
        #** RapidAPI info -> 'description', 'name_translation'
        #** QuranAPI info -> 'id', 'revelation_place', 'revelation_order', 'bismillah_pre'
        #**                  'verses_count', 'name_simple', {'translated_name': 'name'}
        async def _get_rapid_info():
            return await self._get_rapidapi_info(surahID)
        
        async def _get_quranapi_info():
            main_url = self.url.quran_url
            main_endpoint = f'{main_url}/api/v4/'
            base_url = f'{main_endpoint}chapters'
            info_url = f'{base_url}/{surahID}'
            base_response, info_response = await asyncio.gather(
                                            self._request(endpoint=surahID, url=base_url,
                                                            slash=True),
                                            self._request(endpoint='info', url=info_url,
                                                            slash=True)
                                                                )
            base_contents, info_contents = base_response['chapter'], info_response['chapter_info']
            base_keys = ('id', 'name_simple', 'revelation_place', 'revelation_order',
                        'bismillah_pre', 'verses_count')
            info_keys = ('short_text', 'text')
            base_contents = OrderedDict({key: base_contents[key] for key in base_keys})
            info_contents['text'] = [i.text for i in BeautifulSoup(info_contents['text'], 'html.parser')][1:]
            info_contents = OrderedDict({key: info_contents[key] for key in info_keys})
            info_contents['quran-source'] = source
            base_contents.update(info_contents)
            return base_contents
        
        rapidapi_info, quranapi_contents = await asyncio.gather(
                        _get_rapid_info(),
                        _get_quranapi_info()
                        )
        quranapi_contents.update(rapidapi_info)
        return quranapi_contents


    async def extract_rabbana_duas(self, export=False):
        async def _get_dua_contents():
            url = self.url.myislam
            endpoint = '40-rabbana-dua-best-quranic-dua'
            full_page = await self._extract_contents(url=url, endpoint=endpoint, slash=True)
            stripped_page = [i for i in ''.join([i.text for i in full_page]).split('\n') if i]
            start, end = [idx for idx,i in enumerate(stripped_page) if re.match(r'(40 Rabbana Duas)|(Share:)',i)]
            fixed_page = stripped_page[start:end-1]
            return fixed_page

        def _get_indexes(dua_page):
            dua_page.pop(0) #** Removes title
            dua_title_idx = [(idx, i) for idx,i in enumerate(dua_page) if re.match(r'^(Rabbana Dua #\d{1,2})',i)]
            grouped_indexes = [(dua_title_idx[i], dua_title_idx[i+1]) for i in range(0, len(dua_title_idx)-1)]
            grouped_indexes.append((grouped_indexes[-1][-1], (len(dua_page), grouped_indexes[-1][-1][-1])))
            return grouped_indexes

        async def _structure_page():
            dua_page = await _get_dua_contents()
            rabbana_indexes = _get_indexes(dua_page)
            verseID_pattern = r'(\d{1,3}:\d{1,3})'
            verse_pattern = r'\“(.*?)\”'
            dua_contents = OrderedDict({
                                'Rabbana Dua #1': {
                                'verseID': re.search(verseID_pattern, dua_page[1]).group(),
                                'verse-ar': dua_page[0],
                                'verse-en': re.search(verse_pattern, dua_page[1]).group(),
                                'transliteration': dua_page[1][:dua_page[1].find('“')],
                                'Recommended use': dua_page[2].removeprefix('Recommended use:')
                                }})
            for i in rabbana_indexes:
                start, end = i[0][0], i[1][0]
                key = i[0][-1]
                if key not in dua_contents:
                    dua_contents[key] = {}
                contents = dua_page[start:end]
                #** 0: key(Rabbana Dua #\d /Removed), 1: ar_text, 2: verse, 3: Recommended use
                dua_contents[key] = {
                            'verseID': re.search(verseID_pattern, contents[2]).group(),
                            'verse-ar': contents[1],
                            'verse-en': re.search(verse_pattern,contents[2]).group(),
                            'transliteration': contents[2][:contents[2].find('“')],
                            'Recommended use': contents[-1].removeprefix('Recommended use:')}
            return dua_contents

        all_dua_contents = await _structure_page()
        if export:
            return self._exporter(all_dua_contents, '40-Rabbana-Duas', path='jsons/quran/duas')
        return all_dua_contents

    async def extract_all_duas(self, export=False):
        main_url = self.url.islamic_finder
        main_endpoint = 'duas'
        html_contents = await self._extract_contents(url=main_url, endpoint=main_endpoint, slash=True, class_='nav-container stick-on-top')
        all_dua_endpoints = [i['href'] for i in html_contents[0].find_all('a', href=True)][2:]
        ramandan_endpoints, masnoon_endpoints = (all_dua_endpoints[:8], all_dua_endpoints[8:])
        ramandan_endpoints.pop(2) #! `/breaking-fast` endpoint broken
        
        async def _parse_endpoints(key, endpoints):
            dua_categories = OrderedDict({key: {}})
            for idx, endpoint in enumerate(endpoints, start=1):
                response = await self._extract_contents(url=main_url, endpoint=endpoint, slash=True)
                ar_text = response.find_all(class_='large-12 columns arabic')
                transliteration, translation, reference = [response.find_all(class_=i) for i in ('large-12 columns transliteration', 'large-12 columns translation', 'large-12 columns reference')]
                dua_key = endpoint.split('/')[-2].replace('-',' ').title()
                dua_contents = {}
                match len(ar_text):
                    case 1:
                        dua_contents = {dua_key: {
                                                'ar-text': ar_text[0].text[::-1].strip(),
                                                'transliteration': transliteration[0].text,
                                                'translation': translation[0].text,
                                                'reference': reference[0].text}}
                        dua_categories[key][idx] = dua_contents
                    case _ if len(ar_text)>=2:
                        ar_text = {idx: i.text[::-1].strip() for idx,i in enumerate(ar_text, start=1)}
                        transliteration, translation, reference = ({idx: i.text for idx, i in enumerate(items, start=1)} for items in (transliteration, translation, reference))
                        rest_duas = OrderedDict()
                        for idx_, (ar,tr,tra,ref) in enumerate(zip(ar_text, transliteration, translation, reference), start=1):
                            dua_contents = {
                                            'ar-text': ar_text.get(idx_),
                                            'transliteration': transliteration.get(idx_),
                                            'translation': translation.get(idx_),
                                            'reference': reference.get(idx_)}
                            new_dua_key = f'{idx_}-{dua_key}'
                            if new_dua_key not in rest_duas:
                                rest_duas[new_dua_key] = {}
                            rest_duas[new_dua_key] = dua_contents
                        dua_categories[key][idx] = rest_duas
            return dua_categories
        ramadan_duas, masnoon_duas = await asyncio.gather(
                                            _parse_endpoints('Ramadan-Duas', ramandan_endpoints),
                                            _parse_endpoints('Dua-Categories', masnoon_endpoints))
        ramadan_duas.update(masnoon_duas)
        if export:
            self._exporter(ramadan_duas, '142-Duas-with-ramadan', path='jsons/quran/duas')
            return await self._merge_all('all-duas', 'duas')
        return ramadan_duas

class HadithAPI(BaseAPI):
    def __init__(self):
        super().__init__()
        self.url = self.config
    
    async def get_all_hadiths(self, **kwargs):
        '''Fetch all hadiths or specify `author`: str'''
        async def _get_hadiths(**kwargs):
            return await _extract_urls(**kwargs)
        
        async def _extract_urls(**kwargs):
            async def _parser(contents):
                for book, link in tqdm(contents.items(), total=len(contents), desc='Processing Hadiths', colour='green', unit='MB'):
                    with open(self.path / 'jsons' / 'hadiths' / f'book_{book}.json', mode='w', encoding='utf-8') as file2:
                        hadith_json = await self._request(endpoint='', slash=False, url=link)
                        json.dump(hadith_json, file2, indent=4)
                return file2
            default_values = (False, 'English')
            parser, lang = (kwargs.get(key, default_values[i]) for i,key in enumerate(('parser', 'lang')))
            json_file = await self._request(endpoint='', slash=False, url=self.url.hadith_url)
            contents_ = [(nested('book', j, wild=True), nested('link', j, wild=True)) for i in json_file.values() for j in i['collection'] if j.get('language') == lang]
            contents = {key[0][0]: key[1][0] for key in contents_}
            path = Path(deepcopy(self.path))
            file = open(path / 'jsons' / 'hadith_api_links.json', mode='w', encoding='utf-8')
            json.dump(contents, file, indent=4)
            file.close()
            if parser:
                json_file = json.load(open(path / 'jsons' / 'hadith_api_links.json', encoding='utf-8'))
                await _parser(json_file)
                return contents
            else:
                return contents
        
        contents = await _get_hadiths(parser=True)
        book_authors = contents.keys()
        default_values = ['', False]
        author, _  = [kwargs.get(key, default_values[i]) for i, key in enumerate(('author', ''))]
        if author:
            author = self.best_match(author, values_=book_authors)[0]
            book_json = json.loads((self.path / 'jsons' / 'hadiths' / f'book_{author}.json').read_text(encoding='utf-8'))
            return book_json
        else:
            return contents

class IslamFacts(BaseAPI):
    facts = set()
    allah_names = dict()
    
    def __init__(self):
        super().__init__()
        self.url = self.config
    
    @classmethod
    def _update_facts(cls, facts: set):
        file2 = cls._load_file(path=cls.path, name='islam_facts', mode='r', folder='jsons', type_='json')
        fun_facts = dict.fromkeys(file2)
        fun_facts.update(facts)
        cls.facts.update(facts)
        file3 = open(cls.path / 'jsons' / 'islam_facts.json', mode='w', encoding='utf-8')
        json.dump(fun_facts, file3, indent=4, ensure_ascii=False)
        file3.close()
        return fun_facts
    
    async def extract_islam_facts(self, export=False, **kwargs):
        def _randomizer(dict_):
            #?>Modify for more flexibily to show a random content for each method
            new_dict = list(nested('Facts', dict_)[0].values())
            rand_fact = choice(new_dict)
            return rand_fact
        
        async def _extract_gen_facts():
            url = self.url.islam_facts
            endpoint = 'random/facts/religion/islam'
            if len(self.facts) == 0:
                #!> FunFact generator website only allows ~18 free SAME random facts
                while len(self.facts) <= 18:
                    for _ in range(limit):
                        soup = await self._extract_contents(endpoint=endpoint, slash=True, 
                                                            url=url, tag_='h2')
                        fun_fact = soup[0].text
                        formatted = re.sub(r'\((Religion > Islam )\)', '', fun_fact).strip()
                        self.facts.add(formatted)
            set_facts = list(dict.fromkeys(self.facts, {}).keys())
            fun_facts = OrderedDict({'Source': 'https://fungenerators.com',
                        'Facts': OrderedDict({f'Fact-{idx}': fact for idx,fact in enumerate(set_facts, start=1)})})
            return fun_facts

        async def _extract_alislam_facts():
            url = self.url.al_islam
            endpoint = 'articles/25-interesting-facts-you-should-know-about-islam/'
            async def _get_soup():
                soup = await self._extract_contents(url=url,endpoint=endpoint,slash=True, tag_='p')
                fixed_soup = [re.sub(r'(Fact\s{1}\#\d{1,3}:\s{1})','',i) for i in [i.text for i in soup] if i]
                return fixed_soup

            async def _structure_contents():
                contents = await _get_soup()
                all_facts = OrderedDict({'Source': url,
                            'Introduction': contents.pop(0),
                            'Facts':OrderedDict({f'Fact-{idx}': self.add_line_breaks(fact, 11) for idx,fact in enumerate(contents, start=1)})})
                return all_facts
            alislam_facts = await _structure_contents()
            return alislam_facts
        
        limit = kwargs.get('limit', 2)
        if export or not Path(self.path / 'jsons' / 'islamic-facts.json').is_file():
            gen_fun_facts, alislam_facts = await asyncio.gather(
                                                        _extract_gen_facts(),
                                                        _extract_alislam_facts())
            all_facts = DataLoader.add(gen_fun_facts, alislam_facts, key1=1, key2=2)
            self._exporter(all_facts, 'islamic-facts')
            rand_fact = _randomizer(all_facts)
            return rand_fact
        else:
            facts_file = DataLoader.load_file(path='jsons', file_name='islamic-facts')
            self.facts = facts_file
            rand_fact = _randomizer(facts_file)
            return rand_fact
    
    async def extract_allah_contents(self, export=False):
        async def _get_name_contents():
            async def extract_content(**kwargs):
                return await self._extract_contents(**kwargs)
            
            async def _get_ar_names():
                return [i.text[::-1].strip('\n') for i in allah_names_html][1::4]
            
            async def _allah_99names():
                all_names = {idx: key for idx, key in enumerate([i.text for i in main_page], start=1)}
                return all_names
            
            def _fixer(undo=False):
                if not undo:
                    return tuple(map(lambda i: i.translate(''.maketrans('dh',' z')).lstrip() if i.startswith('d') else i, filter_names))
                return names_copied
            
            def _extract_name_data():
                #? Arabic encoding \b[\u0600-\u06FF]+\b
                #?                 \b[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+\b
                summary_ = sorted(''.join([i.text for i in html_contents[3]]).split('\n'), key=len, reverse=True)[0]
                summary = ' '.join([i.strip('();,.')[::-1] if re.findall(r'[();,.]', i) and re.findall(r'\b[\u0600-\u06FF]+\b',i) else i for i in summary_.split()])
                contents = {
                    'transliteration_eng': re.sub(r'[()]', '', html_contents[0][0].text),
                    'transliteration_ar': all_ar_names[ar_name_idx],
                    'description': [i.text for i in html_contents[1]],
                    'mentions-from-quran-hadith': [i.text for i in html_contents[2]],
                    'summary': summary
                }
                modified_contents = {key: ''.join(value) for key,value in contents.items()}
                return modified_contents
            
            main_page, allah_names_html = await asyncio.gather(
                extract_content(endpoint='99-names-of-allah', slash=True, url=self.url.myislam, class_='transliteration'),
                extract_content(endpoint='', slash=False, url=self.url.allah_names, tag_='td', class_='cb-arabic')
            )
            all_en_names, all_ar_names = await asyncio.gather(
                _allah_99names(),
                _get_ar_names()
            )
            self.allah_names = deepcopy(all_en_names)
            filter_names = [i.lower().replace("'", '') for i in all_en_names.values()]
            names_copied = list(all_en_names.values())
            all_names = _fixer(False)
            all_name_contents = {}
            for ar_name_idx, name in enumerate(all_names):
                new_endpoint = f'99-names-of-allah/{name}'
                html_contents = await asyncio.gather(
                                *[extract_content(
                                                endpoint=new_endpoint, slash=True,
                                                url=self.url.myislam, class_=i)
                                for i in ('name-meaning', 'summary', 
                                        'column-section', 'second-section')]
                                )
                org_names = _fixer(True)
                all_name_contents[org_names[ar_name_idx]] = _extract_name_data()
            return all_name_contents
        
        all_contents = await _get_name_contents()
        merged_contents = {}
        for idx, (name, information) in tqdm(enumerate(all_contents.items(), start=1),
                                        total=len(all_contents), desc='Processing Names of Allah',
                                        colour='green', unit='MB'):
            merged_contents[idx] = {
                                    'Name': name,
                                    'Information': {**information}
                                    }
        if export:
            return self._exporter(merged_contents, 'list_allah_names')
        return merged_contents
    
    async def extract_islamic_terms(self, export=False):
        async def _get_soup(query):
            url = self.url.alim
            soup = await self._extract_contents(endpoint=f'islamic-terms-dictionary/{query}', url=url, slash=True)
            return soup
        
        async def _parse_letter(letter):
            letter_dictionary = {}
            letter_words = await _get_soup(letter)
            cleaned_words = ''.join([i.text for i in letter_words]).split('\n')
            all_words = [
                        OrderedDict({
                            'Term': cleaned_words[i],
                            'Definition': ' '.join([i.capitalize() for i in re.split(r'(?<=[.!?])\s+', cleaned_words[i+1])])
                        })
                        for i in range(len(cleaned_words) - 1)
                        if cleaned_words[i].startswith(letter.upper()) and cleaned_words[i + 1]
                    ]
            #** [::] Some had random content
            letter_dictionary[letter.upper()] = None if len(all_words)==0 else \
                                                all_words[7:] if letter=='T' else \
                                                all_words[2:] if letter=='R' else \
                                                all_words
            return letter_dictionary
        
        async def _extract_all():
            queries = ascii_lowercase
            with ThreadPoolExecutor(max_workers=cpu_count()//2) as executor:
                loop = asyncio.get_event_loop()
                tasks = [await loop.run_in_executor(executor, _parse_letter, query) for query in queries]
                islamic_dictionary = await asyncio.gather(*tasks)
            return islamic_dictionary
        
        full_dictionary = await _extract_all()
        if export:
            return self._exporter(full_dictionary, 'islamic-terms')
        return full_dictionary
    
    async def extract_islam_pillars(self, export=False):
        pdf_file = load(path='pdfs', file_name='5-Pillars-of-Islam', ext='pdf')

        def _get_intro():
            '''Summary, Verse, Description'''
            first_page = self._extractor(pdf_file, page_numbers=range(0,1))
            start, middle, third, end = [idx for idx,i in enumerate(first_page) if re.search(r'((^Summary:)|(\“Righteousness)|((Qur\'an 2:177, trans. Abdel Haleem)))|(world\, including in America)',i)]
            summary = ''.join(first_page[start:middle]).removeprefix('Summary: ')
            verse = [i.strip() for i in first_page[middle:third+1]]
            description = first_page[third+1:end+1]
            return summary, verse, description

        def _clean_pdf():
            file = self._extractor(pdf_file)
            copyright = 'Copyright ©2021'
            contact = 'contact the Pluralism Project'
            cleaned_pdf = [i for i in file if not re.search(fr'{copyright}|{contact}',i)]
            return cleaned_pdf

        def _get_pillars():
            cleaned_pdf = _clean_pdf()
            pillars = r'(Shahadah: )|(Salat: )|(Zakat: )|(Sawm: )|(Hajj: )'
            pillars_idx = [(idx, i) for idx,i in enumerate(cleaned_pdf) if re.match(pillars, i)]
            pillar_contents = []
            pillars_idx_grouped = [(pillars_idx[i], pillars_idx[i+1]) for i in range(0, len(pillars_idx)-1)]
            pillars_idx_grouped.append((pillars_idx_grouped[-1][-1], (len(cleaned_pdf), '')))
            for items in pillars_idx_grouped:
                start,end = items[0][0], items[1][0]
                pillar_contents.append(cleaned_pdf[start:end])
            all_pillars = OrderedDict()
            pillar_contents[-1] = pillar_contents[-1][:-1]
            fixed_contents = [''.join(i) for i in pillar_contents]
            for idx,i in enumerate(fixed_contents, start=1):
                pillar_name = re.match(pillars, i).group().replace(': ','')
                contents = re.sub(pillars, '', i)
                key = f'{idx}-{pillar_name}'
                if key not in all_pillars:
                    all_pillars[key] = {}
                readable_contents = self.add_line_breaks(contents, words_per_line=11)
                all_pillars[key] = readable_contents
            return all_pillars

        def _structure_pillars():
            summary, verse, description = _get_intro()
            summary = self.add_line_breaks(summary, 11)
            all_pillars = _get_pillars()
            structured_pillars = OrderedDict({
                                    'Summary': summary,
                                    'Verse': verse,
                                    'Description': description,
                                    '5-Pillars': all_pillars
                                             })
            return structured_pillars

        all_pillar_contents = _structure_pillars()
        if export:
            return self._exporter(all_pillar_contents, 'islamic-pillars-harvard', 'jsons/pillars')
        return all_pillar_contents

    async def extract_all_kalimas(self, export=False):
        url = self.url.muslim_quran
        endpoint = 'kalmas'
        soup = await self._extract_contents(url=url, endpoint=endpoint, slash=True)
        ar_text, en_text = [[i.text for i in soup.find_all(class_=i)] for i in ('arabic-text kalma-arabic noorehudaFont arabic-line-height text-right font_36', 'mt-3')]
        meaning = [i for i in [i.text for i in soup.find_all(class_='col-md-12 pt-4 px-0')] if i]
        names = [i.text for i in soup.find_all(class_='mt-1')]
        descr = [i for i in ''.join([i.text for i in soup.find_all(class_='font_14 override_styles')]).split(('\n'))[:-3] if i and not re.match(r'(\r)|(\xa0)',i)]
        title_ = descr.pop(0)
        title = title_[:title_.find(' Here')]
        all_contents = OrderedDict({'title': self.add_line_breaks(title, 12),
                                   'kalimas': {}})
        kalimas_content = {}
        for idx, (ar, en, trans, name, desc) in enumerate(zip(ar_text, en_text, meaning, names, descr), start=1):
            if idx not in kalimas_content:
                kalimas_content[idx] = {}
            kalimas_content[idx] = OrderedDict({'description': desc,
                                               'ar-text': self.add_line_breaks(ar, 12),
                                               'en-text': self.add_line_breaks(en, 12),
                                               'transliteration': self.add_line_breaks(trans, 11)})
        all_contents['kalimas'] = kalimas_content
        if export:
            return self._exporter(all_contents, 'all-kalimas')
        return all_contents

    async def extract_funeral_guidance(self, export=False):
        pdf_file = load(path='pdfs', file_name='funeral-guidance', ext='pdf')
        pdf_contents = self._extractor(pdf_file)
        first_page = self._extractor(pdf_file, maxpages=1)

        def _get_front_cover():
            title = first_page.pop(0).strip()
            summary = ''.join(first_page[:first_page.index('Introduction  ')]).strip()
            end_intro_idx = [idx for idx,i in enumerate(first_page) if re.search(r'(funeral service)', i)][-1]
            intro_descr = ''.join(first_page[first_page.index('Introduction  '):end_intro_idx+1]).removeprefix('Introduction  ').rstrip(':  ')
            return title, summary, intro_descr

        def _get_toc_contents():
            end_intro_idx = [idx for idx,i in enumerate(first_page) if re.search(r'(funeral service)', i)][-1]
            toc_contents = [i for i in first_page[end_intro_idx+1:][:5]]
            toc_indexes = [(idx, i) for idx, i in enumerate(pdf_contents) if i in toc_contents]
            return pdf_contents
        
        return _get_toc_contents()

        

    @classmethod
    @property
    def allah_99names(cls):
        #!> Add Exception Handling if None
        return cls.allah_names
    
    @classmethod
    @property
    def get_all_facts(cls):
        return cls.facts

class PrayerAPI(BaseAPI):
    def __init__(self):
        super().__init__()
        self.url = self.config
        
    async def extract_qibla_data(self, export=False):
        async def _get_qibla(**kwargs):
            @lru_cache(maxsize=1)
            def _get_coords():
                Coords = namedtuple('Coords', ['lat', 'long', 'qibla'], defaults=['25.4106386', '51.1846025', None])
                loc = location(place)
                coords_qib = Coords(lat=loc.lat, long=loc.lng)
                return coords_qib
            
            place = kwargs.get('place', 'Saudia Arabia')
            coords_qib = _get_coords()
            endpoint = f'qibla/:{coords_qib.lat}/:{coords_qib.long}'
            url = self.url.aladhan
            response = await self._request(endpoint=endpoint, slash=True, url=url)
            qibla_dir = response['data'].get('direction') if response['status']=='OK' else 68.92406695044804
            coords_qib = coords_qib._replace(qibla=qibla_dir)
            return coords_qib
        
        @lru_cache(maxsize=1)
        def _extract_countries():
            pdf_contents = extract_pages(pdf_file)
            countries = ''.join([j.get_text() for i in pdf_contents for j in i])
            all_countries_ = countries.split('\n')[:-1]
            all_countries_.sort(key=lambda i: i[0])
            all_countries = dict.fromkeys(all_countries_, {})
            return all_countries
        
        pdf_file = self.path / 'pdfs' / 'all_countries.pdf'
        all_countries = _extract_countries()
        
        qibla_data = {}
        for idx, (country, _) in tqdm(enumerate(all_countries.items(), start=1), total=len(all_countries), desc='Processing Qibla Data', unit='MB', colour='green'):
            lat, long, qibla_dir = await _get_qibla(place=country)
            qibla_data[idx] = {'Country': country,
                                'latitude': lat,
                                'longitutde': long,
                                'qibla_dir': qibla_dir}
        if export:
            return self._exporter(qibla_data, 'qibla_data')
        else:
            return qibla_data

class ProphetStories(BaseAPI):
    def __init__(self):
        super().__init__()
        self.url = self.config
        
    async def _empty_stories(self):
        async def _get_prophets():
            main_endpoint = 'prophet-stories/'
            soup, stories_soup = await asyncio.gather(
                                self._extract_contents(endpoint=main_endpoint, slash=True,
                                                        url=self.url.myislam, class_='et_pb_text_inner'),
                                self._extract_contents(endpoint=main_endpoint, slash=True,
                                                        url=self.url.myislam)
                                )
            
            def _get_empty():
                def _fix_intros(story, prophet):
                    story = ''.join(story)
                    first_sub = re.sub(r'\d{1,3}\.\s', '', story)
                    second_sub = re.sub(rf'Story of {prophet}', '', first_sub)
                    final_sub = re.sub(r'\xa0', '', second_sub)
                    return final_sub
                
                main_stories = [i.text.split('\n') for i in soup if re.search(r'\d{1,2}\.', i.text)]
                all_prophets_ = [j.removeprefix('Story of ') for i in main_stories for j in i if re.findall(r'Story of', j)]
                all_prophets = {idx: key for idx,key in enumerate(all_prophets_, start=1)}
                prophet_intro = ''.join([i.text for i in soup if re.search(r'(Surah Nahl Ayat)', i.text)]).split('\n')[:-2]
                empty_stories = {'About Prophets': prophet_intro}
                for idx, (story, (_, prophet)) in enumerate(zip(main_stories, all_prophets.items()), start=1):
                    intro_story = _fix_intros(story, prophet)
                    empty_stories[idx] = {prophet: {'Intro': intro_story}}
                return empty_stories
            
            def _prophet_endpoints():
                pattern = re.compile(r".*Story of Prophet.*")
                links = stories_soup.find_all('p')
                prophet_endpoints = [i.a['href'].split('/')[-2] for i in links if re.search(pattern, i.get_text())]
                return prophet_endpoints
            
            return (_prophet_endpoints(), _get_empty())
        
        #** prophet_endpoints, empty_stories
        contents = await _get_prophets()
        return contents
    
    async def _extract_stories(self):
        prophets, empty_stories = await self._empty_stories()
        with ThreadPoolExecutor(max_workers=cpu_count() // 2) as executor:
            loop = asyncio.get_event_loop()
            tasks = [await loop.run_in_executor(executor, self._match_func, prophet) for prophet in prophets]
            
            with tqdm(total=len(tasks), desc="Processing Prophets", colour='green', unit='MB') as pbar:
                async def run_task(task):
                    result = await task
                    pbar.update(1)
                    return result

                completed_stories = await asyncio.gather(*[run_task(task) for task in tasks])

        return completed_stories, empty_stories
    
    async def _match_func(self, prophet):
        #** Generic methods: Propets Yunus, Yasa, Yusuf, Saleh, Sulaiman
        method_map = {
            'prophet-ayyub': self._fix_ayyub,
            'prophet-yunus': self._fix_generics,
            'story-of-prophet-lut': self._fix_lut,
            'prophet-idris': self._fix_idris,
            'prophet-dhul-kifl': self._fix_kifl,
            'prophet-nuh': self._fix_nuh,
            'prophet-al-yasa': self._fix_yasa,
            'prophet-yusuf': self._fix_generics,
            'prophet-saleh-story': self._fix_saleh,
            'story-prophet-sulaiman': self._fix_generics,
            'prophet-adam': self._fix_adam,
        }
        method = method_map.get(prophet, None)
        if method:
            return await method(prophet)
        else:
            return None
    
    @staticmethod
    def _fix_name(prophet):
        return re.findall(r'(?:story-)?(?:of-)?(?:prophet-)?(.+)', prophet)[0].title()
    
    async def _parse_html(self, prophet, soup=False, class_=None):
        response = await self._extract_contents(
                                            endpoint=prophet, slash=True,
                                            url=self.url.myislam,
                                            class_=class_)
        if not soup:
            return response
        else:
            return [i.text for i in response]
    
    async def _fix_generics(self, prophet):
        name = self._fix_name(prophet)
        dict_name = f'Prophet {name}'
        all_contents = dict.fromkeys([dict_name], {})
        soup = await self._parse_html(prophet, soup=True)
        clean_contents_ = ' '.join(soup).split('\n')
        clean_contents = [i.replace('\xa0','') for i in clean_contents_ if i and not re.search(r'Prophet Stories', i)]
        pattern = re.compile(r'The Story (?:Of\s)?Prophet (\w+)(?: \((PBUH)\))?', re.IGNORECASE)
        first_key = pattern.search(''.join(clean_contents)).group()
        first_index = clean_contents.index(first_key)
        end_index = clean_contents.index('Support the site?')
        all_contents[dict_name] = {first_key: clean_contents[first_index:end_index-2][1:]}
        return all_contents
    
    async def _fix_ayyub(self, prophet):
        name = self._fix_name(prophet)
        dict_name = f'Prophet {name}'
        all_contents = dict.fromkeys([dict_name], {})
        soup = await self._parse_html(prophet, soup=False, class_='et_pb_section et_pb_section_1 et_section_regular')
        fam_tree_key = [i.div.strong.get_text() for i in soup][0]
        html_contents = [i.text for i in soup]
        cleaned_contents = [i.replace('\xa0', '') for i in ' '.join(html_contents).split('\n') if i and not re.search(rf'{fam_tree_key}|Back To Prophet Stories', i)]
        verse_mentions_key, verse_section = re.findall(rf'Quranic Verses Mentioning\s\w+', ''.join(html_contents))[0], cleaned_contents.index('Quranic Verses Mentioning Ayyub')
        fam_tree_contents, verse_contents = cleaned_contents[:verse_section], cleaned_contents[verse_section+1:]
        all_contents[dict_name] = {
                            fam_tree_key: fam_tree_contents,
                            verse_mentions_key: verse_contents}
        return all_contents
    
    async def _fix_lut(self, prophet):
        name = self._fix_name(prophet)
        dict_name = f'Prophet {name}'
        all_contents = dict.fromkeys([dict_name], {})
        soup = await self._parse_html(prophet, soup=False)
        html_contents = [i.text for i in soup]
        clean_contents_ = ' '.join(html_contents).split('\n')
        clean_contents = [i.replace('\xa0', '') for i in clean_contents_ if i and not re.search(r'Return To Prophet Stories', i)]
        first_key = 'The Story Of Prophet Lut in Islam'
        first_index = clean_contents.index(first_key)
        second_section_key = 'Quran Verses That Mention The Story Of Prophet Lut'
        second_index = clean_contents.index(second_section_key)
        end_index = clean_contents.index('Support the site?')
        first_section_contents, second_section_contents = clean_contents[first_index:second_index][1:], clean_contents[second_index+1:end_index-2]
        all_contents[dict_name] = {
                            first_key: first_section_contents,
                            second_section_key: second_section_contents
                            }
        
        return all_contents
    
    async def _fix_idris(self, prophet):
        name = self._fix_name(prophet)
        dict_name = f'Prophet {name}'
        all_contents = dict.fromkeys([dict_name], {})
        soup = await self._parse_html(prophet, soup=True)
        clean_contents_ = ' '.join(soup).split('\n')
        clean_contents = [i.replace('\xa0', '') for i in clean_contents_ if i and not re.search(r'List Of Prophets In Islam', i)]
        first_key = 'Story of Prophet Idris In Islam'
        first_index = clean_contents.index(first_key)
        second_key = 'Prophet Idris Story'
        second_index = clean_contents.index(second_key)
        end_index = clean_contents.index('Support the site?')
        first_contents, second_contents = clean_contents[first_index:second_index][1:], clean_contents[second_index+1:end_index-2]
        all_contents[dict_name] = {
                        first_key: first_contents,
                        second_key: second_contents
        }
        return all_contents
    
    async def _fix_kifl(self, prophet):
        name = self._fix_name(prophet)
        dict_name = f'Prophet {name}'
        all_contents = dict.fromkeys([dict_name], {})
        soup = await self._parse_html(prophet, soup=True)
        clean_contents_ = ' '.join(soup).split('\n')
        clean_contents = [i.replace('\xa0','') for i in clean_contents_ if i and not re.search(r'List Of Prophets In Islam', i) and not i==' ']
        first_key = 'Prophet Dhul Kifl'
        first_index = clean_contents.index('Prophet Dhul Kifl')
        second_key = 'Story Of Prophet Dhul-Kifl by Ibn jarir'
        second_index = clean_contents.index(second_key)
        end_index = clean_contents.index('Support the site?')
        first_contents, second_contents = clean_contents[first_index:second_index][2:], clean_contents[second_index+1:end_index-1]
        all_contents[dict_name] = {
                        first_key: first_contents,
                        second_key: second_contents
        }
        return all_contents
    
    async def _fix_nuh(self, prophet):
        name = self._fix_name(prophet)
        dict_name = f'Prophet {name}'
        all_contents = dict.fromkeys([dict_name], {})
        soup = await self._parse_html(prophet, soup=True)
        clean_contents_ = ' '.join(soup).split('\n')
        clean_contents = [i.replace('\xa0','') for i in clean_contents_ if i and not re.search(r'Prophet Stories', i)]
        first_key = re.search(rf'Prophet {name}', ''.join(clean_contents)).group()
        first_index = clean_contents.index(first_key)
        second_key = 'References that refer to Prophet Noah'
        second_index = clean_contents.index('Imran (3:33)')
        end_index = clean_contents.index('Support the site?')
        first_contents, second_contents = clean_contents[first_index:second_index-1][1:-1], clean_contents[second_index-1:end_index-2]
        all_contents[dict_name] = {
                    first_key: first_contents,
                    second_key: second_contents
        }
        return all_contents
    
    async def _fix_adam(self, prophet):
        name = self._fix_name(prophet)
        dict_name = f'Prophet {name}'
        all_contents = dict.fromkeys([dict_name], {})
        soup = await self._parse_html(prophet, soup=True, class_='et_pb_module et_pb_code et_pb_code_0')
        clean_contents_ = ''.join(soup).split('\n')
        clean_contents = [i.replace('\xa0','') for i in clean_contents_ if i]
        first_key = 'STORY OF PROPHET ADAM (AS) IN ISLAM'.title()
        second_key = 'Adam (PBUH) learns the names of everything:'
        second_index = clean_contents.index(second_key)
        first_contents, second_contents = clean_contents[:second_index], clean_contents[second_index+1:]
        all_contents[dict_name] = {
                    first_key: first_contents,
                    second_key.strip(':'): second_contents
        }
        return all_contents
    
    async def _fix_yasa(self, prophet):
        name = self._fix_name(prophet)
        dict_name = f'Prophet {name}'
        all_contents = dict.fromkeys([dict_name], {})
        soup = await self._parse_html(prophet, soup=True)
        clean_contents_ = ' '.join(soup).split('\n')
        clean_contents = [i.replace('\xa0','') for i in clean_contents_ if i and not re.search(r'Prophet Stories', i)]
        first_key = re.search(r'Story of Al-Yasa \(Elisha\)', ''.join(clean_contents)).group()
        first_index = clean_contents.index(first_key)
        end_index = clean_contents.index('Support the site?')
        all_contents[dict_name] = {first_key: clean_contents[first_index:end_index-2]}
        return all_contents
    
    async def _fix_saleh(self, prophet):
        name = self._fix_name(prophet)
        dict_name = f'Prophet {name.rstrip("-Story")}'
        all_contents = dict.fromkeys([dict_name], {})
        soup = await self._parse_html(prophet, soup=True)
        clean_contents_ = ' '.join(soup).split('\n')
        clean_contents = [i.replace('\xa0','') for i in clean_contents_ if i and not re.search(r'Prophet Stories', i)]
        pattern = re.compile(r'Story (?:Of\s)?Prophet (\w+)', re.IGNORECASE)
        first_key = pattern.search(''.join(clean_contents)).group()
        first_index = clean_contents.index(first_key)
        end_index = clean_contents.index('Support the site?')
        all_contents[dict_name] = {first_key: clean_contents[first_index:end_index-2][1:]}
        return all_contents
    
    async def extract_all_prophets(self, export=False):
        completed_stories, empty_stories = await self._extract_stories()
        all_contents = {idx: {j: {'Full Story': k, 'Intro': empty_stories[idx].get(j)['Intro']}} 
                            for idx, name in enumerate(completed_stories, start=1) 
                            for j, k in name.items()}

        if export:
            return self._exporter(all_contents, 'prophet_stories')
        return all_contents

class ProphetMuhammad(BaseAPI):
    def __init__(self):
        self.url = self.config
        self.name = 'The Life of the Prophet Muhammad (Peace and blessings of Allah be upon him)'
        self.title = 'The-Life-of-The-Prophet-Muhammad'
    
    async def _pdf_parser(self):
        def _cleaner(contents):
        #** Removes page numbers and header: 'The Life of the Prophet Muhammad (Peace and blessings of Allah be upon him)'
            header_pattern = re.escape(self.name)
            pattern = fr'\d{{1,2}}|{header_pattern}'
            fixed_contents = [re.sub(pattern, '', i) if re.match(pattern, i) else i for i in contents]
            cleaned_contents_ = ' '.join([i.strip() for i in fixed_contents if i]).strip()
            cleaned_contents = '{}{}'.format(cleaned_contents_, '.' if not cleaned_contents_.endswith('.') else '')
            return cleaned_contents
        
        async def _title_page():
            def _get_toc():
                #** Table of Contents
                toc_page = self._extractor(pdf, page_numbers=[3])
                toc_contents = [i.split('.', 1)[0].rstrip() for i in toc_page[3:-1]]
                updated_toc = {idx: content for idx, content in enumerate(toc_contents, start=1)}
                return updated_toc
            
            title_page = self._extractor(pdf, maxpages=2)
            title_contents = list(map(lambda i: i.strip(), title_page))[:-8]
            title = ' '.join([title_contents.pop(0) for _ in range(3)])
            translation_ar, translation_en = [title_contents.pop(0) for _ in range(2)]
            toc_contents = _get_toc()
            updated_title = {title: {
                                    'introduction': ''.join(title_contents),
                                    'transliteration_ar': translation_ar,
                                    'transliteration_en': translation_en,
                                    'table_of_contents': {**toc_contents}
                                    }}
            return updated_title
        
        async def _parse_chap(start_page, end_page):
            chap_page = self._extractor(pdf, page_numbers=range(start_page, end_page + 1))[3:]
            chap_contents = _cleaner(chap_page)
            return chap_contents
        
        async def _parse_gloss():
            def _clean_gloss():
                glossary_page = self._extractor(pdf, page_numbers=range(85, 91))[3:]
                header_pattern = re.escape(self.name)
                pattern = fr'{header_pattern}'
                glossary_contents = [i.strip() for i in glossary_page if not re.match(pattern, i)]
                cleaned_contents = [i for i in glossary_contents if i and not re.match(r'\d{2}', i)]
                return cleaned_contents
            
            gloss_contents = _clean_gloss()
            abd_allah = gloss_contents[0], ' '.join(gloss_contents[1:3])
            abd_ibn_1 = gloss_contents[3].split('   ')[0]+'-1', gloss_contents[3].split('   ')[1]
            ubayy = gloss_contents[4], ' '.join(gloss_contents[5:8])
            abd_al = gloss_contents[8], gloss_contents[9]
            muttalib = gloss_contents[10], ' '.join(gloss_contents[11:13])
            abd_ibn_2 = gloss_contents[13].split('   ')[0]+'-2', gloss_contents[13].split('   ')[1]
            abu_rabiah = gloss_contents[14].split('   ')[0], gloss_contents[14].split('   ')[1]
            abdu_manaf = gloss_contents[15].split('   ')[0], gloss_contents[15].split('   ')[1]+gloss_contents[16]
            abrahah = gloss_contents[17], ' '.join(gloss_contents[18:20])
            abraham_ib = ' '.join(gloss_contents[20:22]), ' '.join(gloss_contents[22:29])
            abo_bakr = gloss_contents[29], ' '.join(gloss_contents[30:34])
            abu_dujanah = gloss_contents[34].split('   ')[0], gloss_contents[34].split('   ')[1]+ ' '.join(gloss_contents[35:37])
            abujahl = gloss_contents[37], ' '.join(gloss_contents[39:46])
            abu_sufyan = gloss_contents[38], ' '.join(gloss_contents[46:50])
            abo_talib = gloss_contents[50], ' '.join(gloss_contents[51:55])
            addas = gloss_contents[55], ' '.join(gloss_contents[56:60])
            adhan = gloss_contents[60], gloss_contents[61]
            aisah = gloss_contents[62], ' '.join(gloss_contents[63:65])
            al_abbas = gloss_contents[65], ' '.join(gloss_contents[66: 69])
            ali = gloss_contents[69], ' '.join(gloss_contents[70:73])
            allahu_akar = gloss_contents[73].split('   ')[0], gloss_contents[73].split('   ')[1]
            alms = gloss_contents[74], gloss_contents[75]
            aminah = gloss_contents[76], gloss_contents[77]
            amro_ibun_ = gloss_contents[78].split('   ')
            amro_ibun = f'{amro_ibun_[0]} {amro_ibun_[1]}', ' '.join(amro_ibun_[2:])
            al_ass = gloss_contents[79], ' '.join(gloss_contents[81:84])
            ansar = gloss_contents[80], '{} {} {}'.format(' '.join(gloss_contents[84:86]), gloss_contents[87], gloss_contents[86])
            apostle = gloss_contents[88], gloss_contents[89]
            
            #!> FINISH GLOSSARY
            updated_gloss = OrderedDict({
                            abd_allah[0]: abd_allah[1],
                            abd_ibn_1[0]: abd_ibn_1[1],
                            ubayy[0]: ubayy[1],
                            abd_al[0]: abd_al[1],
                            muttalib[0]: muttalib[1],
                            abd_ibn_2[0]: abd_ibn_2[1],
                            abu_rabiah[0]: abu_rabiah[1],
                            abdu_manaf[0]: abdu_manaf[1],
                            abrahah[0]: abrahah[1],
                            abraham_ib[0]: abraham_ib[1],
                            abo_bakr[0]: abo_bakr[1],
                            abu_dujanah[0]: abu_dujanah[1],
                            abujahl[0]: abujahl[1],
                            abu_sufyan[0]: abu_sufyan[1],
                            abo_talib[0]: abo_talib[1],
                            addas[0]: addas[1],
                            adhan[0]: adhan[1],
                            aisah[0]: aisah[1],
                            al_abbas[0]: al_abbas[1],
                            ali[0]: ali[1],
                            allahu_akar[0]: allahu_akar[1],
                            alms[0]: alms[1],
                            aminah[0]: aminah[1],
                            amro_ibun[0]: amro_ibun[1],
                            al_ass[0]: al_ass[1],
                            ansar[0]: ansar[1],
                            apostle[0]: apostle[1]
                            })
            return updated_gloss
        
        pdf = self._get_file(path=self.path, file_name=self.title)
        title_page, glossary = await asyncio.gather(
                                    _title_page(),
                                    _parse_gloss()
                                    )
        
        chapters = {
                    1: (4, 8),    2: (8, 10),   3: (10, 11), 
                    4: (12, 13),  5: (14, 15),  6: (16, 17),
                    7: (18, 19),  8: (20, 22),  9: (23, 24),
                    10: (25, 27), 11: (28, 29), 12: (30, 32),
                    13: (33, 35), 14: (36, 37), 15: (38, 39),
                    16: (40, 41), 17: (42, 44), 18: (45, 46),
                    19: (47, 50), 20: (51, 54), 21: (55, 59),
                    22: (60, 63), 23: (64, 68), 24: (69, 72),
                    25: (73, 75), 26: (76, 78), 27: (79, 82),
                    28: (83, 84)
                }
        
        with ThreadPoolExecutor(max_workers=cpu_count() // 2) as executor:
            loop = asyncio.get_event_loop()
            tasks = [await loop.run_in_executor(executor, _parse_chap, *chap) for chap in chapters.values()]
            chapters = await asyncio.gather(*tasks)
        chapters = {idx: contents for (idx, contents) in enumerate(chapters, start=1)}
        return (title_page, chapters, glossary)
    
    async def proph_muhammads_life(self, export=False):
        title_page, chapters, glossary = await self._pdf_parser()
        toc_contents = nested('table_of_contents', title_page)[0]
        toc_and_chaps = {idx: {chap: chapters.get(idx) if idx!=29 else glossary}
                            for idx, chap in tqdm(enumerate(toc_contents.values(), start=1),
                                                total=len(toc_contents.values()),
                                                desc='Processing Life-of-Prophet-Muhammad PDF',
                                                unit='MB',
                                                colour='green')}
        title_page[self.name]['table_of_contents'] = toc_and_chaps
        full_book = deepcopy(title_page)
        
        if export:
            return self._exporter(full_book, 'life_of_prophet_muhammad')
        return full_book

    async def extract_muhammads_names(self, export=False):
        url = self.url.urdu_point
        endpoint = 'islam/99names_Allah.php'
        soup = await self._extract_contents(url=url, endpoint=endpoint, slash=True)
        
        def _get_title():
            title_soup = soup.find_all(class_='phead')
            title = title_soup[0].text
            return title

        def _get_table():
            table_soup = soup.find_all(class_='resp_table')
            fixed_table = [i for i in ''.join([i.text for i in table_soup]).split('\n') if i and not re.match(r'^\d', i)][4:]
            grouped_table = [fixed_table[i:i+3] for i in range(0, len(fixed_table)-1, 3)]
            return grouped_table

        def _structured_table():
            title, table = _get_title(), _get_table()
            full_table = OrderedDict({'Source': url,
                                    f'{title}': {}})
            table_contents = {}
            for idx, (ar_name, trans, meaning) in enumerate(table, start=1):
                if idx not in table_contents:
                    table_contents[idx] = {}
                table_contents[idx] = OrderedDict({'Name': ar_name,
                                                    'Transliteration': trans,
                                                    'Meaning': meaning})
            full_table[f'{title}'] = table_contents
            return full_table

        all_names = _structured_table()
        if export:
            return self._exporter(all_names, 'list_of_prophet_muhammed_names')
        return all_names

class IslamicStudies(BaseAPI):
    def __init__(self):
        super().__init__()
        self.url = self.config
        
    async def islamic_timeline(self, export=False):
        async def _get_timeline():
            timeline_data = [i.text for i in soup.find_all('p')]
            centuries = [(idx, i) for idx, i in enumerate(timeline_data) if re.search(r'\d{1,2}(?:th) Century', i)]
            return (centuries, timeline_data)
        
        def _get_contents(contents, start, end):
            contents = contents[start+1:end]
            updated_contents = [re.sub(r'\s+',' ',i).replace('\n','').strip() for i in contents][:-1]
            fix623 = [re.sub(r'(:)', r' -', i) if re.match(r'^(623:)', i) else i for i in updated_contents]
            return fix623
        
        async def _get_credits():
            return soup.find('h3').text
            
        async def _parse_timeline():
            time_data, credentials = await asyncio.gather(
                                                    _get_timeline(),
                                                    _get_credits()
                                                    )
            centuries, timeline_data = time_data
            grouped_cent = [(centuries[i], centuries[i + 1]) for i in range(0, len(centuries)-1, 1)]
            all_contents = OrderedDict()
            for _, idx_century in tqdm(enumerate(grouped_cent), total=len(grouped_cent), desc='Processing Islamic Timeline', colour='green', unit='MB'):
                key = idx_century[0][1].replace('\n', '').strip()
                indexes = (idx_century[0][0], idx_century[1][0])
                all_contents[key] = _get_contents(timeline_data, indexes[0], indexes[1])
            all_contents['Credits'] = credentials
            
            if export:
                return self._exporter(all_contents, 'islamic-timeline')
            return all_contents
        
        html_file = open(self.path / 'htmls' / 'Timeline (History of Islam).html', mode='r', encoding='utf-8').read()
        soup = BeautifulSoup(html_file, 'html.parser')
        return await _parse_timeline()
    
    async def get_islam_laws(self, export=False):
        
        @lru_cache(maxsize=1)
        async def _get_docx():
            return Document(self.path / 'docxs' / 'islam-laws.docx')
        
        def _table_extractor(table_index=0, columns=False):
            if table_index < len(docx.tables):
                table = docx.tables[table_index]
                table_data = []
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = ''
                        for paragraph in cell.paragraphs:
                            cell_text += paragraph.text.strip() + '\n'
                        row_data.append(cell_text.strip())
                    table_data.append(row_data)
                if columns:
                    return _table_extractor()[0]
                return table_data
            else:
                return None
        
        def _content_extractor():
            all_contents = {}
            for i in tqdm(range(2), desc='Processing Islamic Laws', unit='MB', colour='green'):
                for j in _table_extractor(i)[1:]:
                    j = list(map(lambda i: i.replace('\n',''), j))
                    all_contents[f'Category-({j[0]})'] = {
                                            'Prohibited (Haram)': j[1],
                                            'Lawful (Halal)': j[2],
                                            'Comments': j[3]
                                            }
            all_contents = {idx: content for idx, content in enumerate(all_contents.items(), start=1)}
            return all_contents
        
        docx = await _get_docx()
        structured_contents = _content_extractor()
        if export:
            return self._exporter(structured_contents, 'islamic-laws')
        return structured_contents
    
    async def road_peace_html(self, export=False):
        
        @lru_cache(maxsize=1)
        async def _get_html():
            html_file = open(self.path / 'htmls' / 'The Road to Peace and Salvation.html', mode='r', encoding='utf-8')
            soup = await self._extract_contents(html_file=html_file)
            return soup
        
        def _get_contents(*args):
            contents, start, end = args
            contents = ''.join(contents[start+1:end]).strip()
            return contents
        
        async def _cleaned_html():
            html = await _get_html()
            html_contents = ''.join([i.text for i in [i for i in html if i]]).split('\n')
            old_html_ = [i for i in html_contents if not re.match(r'(?:\d{1,2})\/16The Road to Peace and Salvation', i)]
            old_html = [i for i in old_html_[1:] if i]
            credentials = html.find('h3').text
            title = ' '.join([old_html.pop(1) for _ in range(3)])
            _ = [old_html.pop(idx) for idx, i in enumerate(old_html) if re.match(r'Chapter (?:\d{1}) ', i)]
            toc_contents = [i[i.find('. ')+2:i.find(' -')] for i in html_contents if re.search(r'(?:\d{1}\.)\s(?:\w+)', i)][:-1]
            intro_ = old_html[old_html.index('INTRODUCTION')+1:old_html.index('ON THE EXISTENCE OF THE DIVINE BEING')]
            intro = ''.join(intro_).strip('Chapter 1 ')
            cleaned_html_ = [re.sub(r'(Chapter (?:\d{1}) )|(?:\d{1,2})\/16', '', i) if re.search(r'(Chapter (?:\d{1}) )|(?:\d{1,2})\/16', i) else i for i in old_html]
            cleaned_toc = {title: 
                        {'Intro': intro,
                        'table_of_contents':
                        {idx: index_name for idx, index_name in enumerate(toc_contents, start=1)},
                        'Credits': credentials}}
            cleaned_html = [i for i in cleaned_html_ if i][8:] #** Excluding toc contents
            return (cleaned_toc, cleaned_html)
        
        async def _structured_contents():
            toc_contents, html = await _cleaned_html()
            title = list(toc_contents.keys())[0]
            
            def _get_toc():
                toc_indexes = [(idx, i) for idx, i in enumerate(html) if i.isupper()]
                toc_grouped = [(toc_indexes[i], toc_indexes[i + 1]) for i in range(0, len(toc_indexes) - 1)]
                return toc_grouped
            
            def _updated_toc():
                toc_indexes = _get_toc()
                updated_toc = {}
                last_idx = []
                for idx, content in tqdm(enumerate(toc_indexes), total=5,
                                        desc='Processing Road-to-Peace-and-Salvation HTML',
                                        colour='green', unit='MB'):
                    if idx==0:
                        #** Intro content already implemented
                        continue
                    key = content[0][1].title()
                    indexes = (content[0][0], content[1][0])
                    contents = _get_contents(html, *indexes)
                    updated_toc[idx] = {key: self.add_line_breaks(contents, 10)}
                    last_indexes = (content[1][0], content[1][1])
                    last_idx.append(last_indexes)
                    if idx==4:
                        key = last_idx[-1][1].title().strip()
                        indexes = (last_idx[-1][0], len(html))
                        contents = _get_contents(html, *indexes)
                        updated_toc[5] = {key: self.add_line_breaks(contents, 10)}
                return updated_toc
            
            all_contents = OrderedDict({
                                    title: {
                                        'Intro': toc_contents[title]['Intro'],
                                        'table_of_contents': _updated_toc(),
                                        'Credits': toc_contents[title]['Credits']
                                            }
                                        })
            return all_contents
        full_contents = await _structured_contents()
        if export:
            return self._exporter(full_contents, 'road_to_peace_salvation')
        return full_contents

    async def extract_life_and_death(self, export=False):
        pdf_file = load(path='pdfs', file_name='life-and-death', ext='pdf')
        pdf_contents = [i for i in self._extractor(pdf_file) if i and not re.match(r'(\d{1,2})', i)]

        def _get_title():
            contents = self._extractor(pdf_file, maxpages=1)[:3]
            title = contents.pop(0)
            author = contents.pop(-1)
            return title, author
        
        def _get_toc():
            toc_contents = self._extractor(pdf_file, page_numbers=range(2,4))[3:-1]
            cleaned_contents = [re.sub(r'(\_+\s{1}\d{1,2})','',i).rstrip() for i in toc_contents]
            return cleaned_contents

        def _group_toc():
            toc = _get_toc()
            toc_indexes = [(idx, i) for idx, i in enumerate(pdf_contents) if i in toc][3:]
            toc_grouped = [(toc_indexes[i], toc_indexes[i+1]) for i in range(len(toc_indexes)-1)]
            toc_grouped.append((toc_grouped[-1][-1], (len(pdf_contents), toc_grouped[-1][-1][-1])))
            return toc_grouped
        
        def _parse_chaps():
            full_toc = _group_toc()
            parsed_chaps = OrderedDict()
            chap_num = None
            for idx, items in enumerate(full_toc, start=1):
                start, end = items[0][0], items[1][0]
                first_chap, second_chap = items[0][-1], items[1][-1]

                if re.search(r'^(Chapter)', first_chap):
                    chap_num = first_chap
                    parsed_chaps[chap_num] = OrderedDict()
                if chap_num is not None:
                    chap_contents = pdf_contents[start+1:end]
                    cleaned_contents = [i for i in chap_contents if not re.match(r'(\s+?\d{1,2})|(\-{0,3}\s+?)|(\‘Every )', i)]
                    parsed_chaps[chap_num][first_chap] = cleaned_contents
            return parsed_chaps

        def _combined():
            front_cover, chapters = _get_title(), _parse_chaps()
            title, author = front_cover[0], front_cover[-1]
            full_book = OrderedDict({'Title': title,
                                    'Author': author,
                                    'Chapters': chapters})
            return full_book
        
        parsed_pdf = _combined()
        if export:
            return self._exporter(parsed_pdf, 'life-and-death')
        return parsed_pdf

class IslamPrayer(BaseAPI):
    def __init__(self):
        self.file = self._get_file(path=self.path, file_name='salah-guide')
    
    async def get_wudu(self, export=False):
        def _cleaner(page_contents):
            return [i for i in page_contents if re.search(r'\w+', i) and not re.match(r'(?:\d{1,2})|(?:How to Perform Wudu\’ \(Step-by-Step\))', i)]
        
        async def _wudu_foundations():
            page = self._get_page(wudu_guide, 10, 11)[:-1]
            title_contents = page.pop(0)
            indexes = self._get_indexes(page, found=True)[0]
            foundation = {title_contents: {}}
            for idx, (i,j) in enumerate(indexes, start=1):
                start, end = i[0], j[0] if j[0] != len(page) else len(page)
                key = i[1].title()
                contents = _get_contents(page, start, end)
                foundation[title_contents][key] = list(contents.split('  '))
                if idx==3:
                    structured_dict = {}
                    current_key = None
                    contents = list(contents.split()[:-1])[:-2]
                    for item in contents:
                        if re.match(r'(?:\d{1})\.', item):
                            current_key = item.strip('.')
                            structured_dict[current_key] = ''
                        elif current_key is not None:
                            structured_dict[current_key] += f'{item} '
                    foundation[title_contents][key] = structured_dict
            return foundation
        
        async def _title_contents():
            page_ = self._get_page(wudu_guide, 11, 12)
            page = _cleaner(page_)
            wudu_title_contents = [page.pop(0).rstrip() for _ in range(3)][::2]
            return wudu_title_contents
        
        def _get_contents(*args):
            contents, start, end = args
            old_contents = ' '.join(contents[start+1:end]).rstrip()
            return old_contents
        
        def _format_page(page, g_index):
            '''g_index: grouped_indexes'''
            steps = {}
            for i, j in g_index:
                start, end = i[0], j[0] if j[0] != len(page) else len(page)
                step_name = i[1]
                contents = _get_contents(page, start, end)
                steps[step_name] = contents
            return steps
        
        async def _wudu_contents():
            pages_ = [self._get_page(wudu_guide, *i) for _,i in enumerate(((11,14), (14, 15)))]
            pages = [_cleaner(i) for i in pages_]
            first_six, nine_ten = pages
            step_sev = first_six.index('Step 7 - Head')
            
            def _sev_eight():
                sev = first_six[step_sev:len(first_six)]
                sev_key, eight_key = [sev.pop(0) for _ in range(2)]
                sev_cont, eight_cont = sev[:4], sev[4:]
                sev_eight = {
                            sev_key: sev_cont,
                            eight_key: eight_cont
                            }
                updated_both = {idx: {key:value} for idx, (key,value) in enumerate(sev_eight.items(), start=7)}
                return updated_both
            
            def _fix_ten():
                label = 'Step 10 - Closing Du’a/Invocation'
                ten = both_fixed[-1].get(label)
                ten_idx = ten.find(label)
                ten_cont = ten[ten_idx+len(label)+1:]
                both_fixed[-1][label] = ten_cont
                both_fixed[-1] = {idx: {key:list(value.split('  '))} for idx,(key,value) in enumerate(both_fixed[-1].items(), start=9)}
                return both_fixed[-1]
            
            updated_six = first_six[3:step_sev]
            merged_pages = updated_six + nine_ten
            sev_eight = _sev_eight()
            both_indexes = [self._get_indexes(i, r'Step \d{1,2}')[0] for _,i in enumerate((updated_six, nine_ten))]
            both_fixed = [_format_page(merged_pages, i) for i in both_indexes]
            both_fixed[-1] = _fix_ten()
            both_fixed[0] = {idx: {key: list(value.split('  '))} for idx,(key,value) in enumerate(both_fixed[0].items(), start=1)}
            both_fixed[0].update(sev_eight)
            both_fixed[0].update(both_fixed[-1])
            wudu_contents = OrderedDict(both_fixed[0])
            return wudu_contents
        
        async def _wudu_rules():
            page = self._get_page(wudu_guide, 15, 16)[:-3]
            title = ''.join([page.pop(0) for _ in range(2)])
            page[page.index('MUSLIM')] = 'NOTE'
            rules = page[:12]
            
            def _fix_rules():
                tip1 = page[12:17]
                tip2 = page[17:21]
                proph_message = page[21:page.index('NOTE')]
                merged_rules = [tip1, tip2, proph_message]
                rule_contents = {contents[0] if re.match(r'(?:Wudu Tip - \w+)', contents[0]) \
                                else ''.join(contents[:2]): \
                                contents[1:] if re.search(r'(?:Wudu Tip - \w+)', contents[0]) \
                                else ''.join(contents[2:]) for contents in merged_rules}
                return rule_contents
                
            rule_contents = _fix_rules()
            note = page[page.index('NOTE'):]
            structured = OrderedDict({
                        title: {
                            'Rules':rules,
                            note[0]: note[1:],
                            'Tips': rule_contents
                            }})
            return structured
        
        
        wudu_guide = self.file
        title_contents, wudu_contents, foundation, rules = await asyncio.gather(
                                                                _title_contents(),
                                                                _wudu_contents(),
                                                                _wudu_foundations(),
                                                                _wudu_rules()
                                                                )
        
        def _merge_all():
            title, desc = title_contents
            all_merged = {
                'Wudu-Guide': {
                            'Introduction': {**foundation},
                            'Mandatory-Rules': {**rules},
                            title: {desc: {**wudu_contents}}
                            }
                        }
            return all_merged
        
        all_contents = _merge_all()
        if export:
            return self._exporter(all_contents, 'wudu-guide')
        return all_contents

    async def get_prayer(self, export=False):
        def _get_contents(*args):
            contents, start, end = args
            old_contents = contents[start+1:end]
            return old_contents
        
        # async def _prayer_rules():
        #     prereqs = self._get_page(prayer_guide, 16, 17)[:-3]
        #     title = prereqs.pop(0)
        #     pre_indexes = self._get_indexes(prereqs, r'(?:\d{1})\. \w+')[0]
        #     prayer_contents = {title: {}}
        #     for idx, (i,j) in enumerate(pre_indexes, start=1):
        #         start, end = i[0], j[0] if j[0] != len(prereqs) else len(prereqs)
        #         key = i[1].title()[3:]
        #         contents = _get_contents(prereqs, start, end)
        #         prayer_contents[title][idx] = {key:contents}
        #     return prereqs
        
        async def _prayer_steps():
            first_rakah = self._get_page(prayer_guide, 17, 24)
            return first_rakah
        
        prayer_guide = self.file
        return await _prayer_steps()

class ArabicAPI(BaseAPI):
    def __init__(self):
        super().__init__()
        self.url = self.config
    
    async def arabic_alphabet(self, export=False):
        path = self.path
        pdf_file = self._get_file(path, 'arabic-alphabet')
        return pdf_file


async def main():
    # tracemalloc.start()

    # a = QuranAPI()
    # b = HadithAPI()
    c = IslamFacts()
    # d = PrayerAPI()
    # e = ProphetStories()
    # f = ProphetMuhammad()
    # g = IslamicStudies()
    # h = IslamPrayer()
    # i = ArabicAPI()
    
    async def run_all(default):
        tasks = [asyncio.create_task(task) for task in [
                    # d.extract_qibla_data(default),
                    # a.surahquran_extract_surahs(default),
                    # a.altafsir_extract_surahs(default),
                    # a.extract_surahs_info(default),
                    # a._surah_base_info(),
                    # b.get_all_hadiths(parser=default),
                    # c.extract_allah_contents(default),
                    # c.extract_islam_facts(default, limit=18),
                    c.extract_funeral_guidance(default),
                    # c.extract_all_kalimas(default),
                    # c.extract_islam_pillars(default)
                    # e.extract_all_prophets(default),
                    # f.proph_muhammads_life(default),
                    # f.extract_muhammads_names(default)
                    # g.islamic_timeline(default),
                    # g.get_islam_laws(default),
                    # g.road_peace_html(default),
                    # g.extract_life_and_death(default)
                    # c.islamic_terms(default),
                    # h.get_wudu(default),
                    # h.get_prayer(default),
                    # a.get_keyword(keyword='woman'),
                    # i.arabic_alphabet(default),
                    # a.extract_all_duas(default),
                    ]]
        results = await asyncio.gather(*tasks)
        return results
    start = time()
    # try:
    #     results = await run_all()
    #     # pprint(results)
    # except Exception as e:
    #     traceback = tracemalloc.get_object_traceback(e)
    #     print(traceback)
    results = await run_all(False)
    end = time()
    pprint(results)
    timer = (end-start)
    minutes, seconds = divmod(timer, 60) 
    print(f"Execution Time: {minutes:.0f} minutes and {seconds:.5f} seconds")

if __name__ == '__main__':
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print('Terminated')






















